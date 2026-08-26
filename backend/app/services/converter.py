"""
Word (.docx) to plain text converter for LMS assessment documents.

Three known template structures are handled automatically based on
document shape — no filename hints required.

  Template A  —  "Assessor Guide" (Learning Vault single-table format)
                  One large table contains both a boilerplate header section
                  (Instructions / Range and conditions / Rubric rows) AND all
                  question content packed into a single merged question cell.
                  Correct answers are indicated by RED font colour.
                  Files: SIRXCEG*, SITEEVT*, SITHACS*, SITHCCC028,
                         SITHFAB023*, SITHFAB025*, SITEEVT029

  Template B  —  Per-question Short-Answer table (SITHCCC / SITX style)
                  A block of body-level metadata paragraphs (title, student
                  name, etc.) followed by one 3-column table per question.
                  The last row of each table is a merged cell whose text
                  already starts with "ASSESSOR KEY: " in dark-green colour.
                  Files: SITHCCC027, SITHCCC029, SITHCCC036,
                         SITHKOP010, SITXFSA005

  Template C  —  Per-question Multiple-Choice table (SITHCCC / SITX style)
                  Same header block as Template B but each table has 4 rows
                  × 3 effective columns.  Options are listed two per row
                  (Col 1 / Col 2).  The correct answer sits in the last row,
                  Col 2, as "ASSESSOR KEY: X" in dark-green colour.
                  Files: SITHCCC031, SITHCCC042, SITHCCC043,
                         SITHPAT016, SITXWHS005

Templates B and C share one extraction path because their "ASSESSOR KEY:"
prefix is already embedded as literal cell text — colour detection is
unnecessary.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
HOW TO ADD A NEW TEMPLATE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. Add a detection branch in _detect_template() that returns a new label
   (e.g. "D") based on a reliable structural signal (table shape, first-cell
   text, presence of body paragraphs, etc.).
2. Implement _extract_template_d(doc) → list[dict] following the same
   pattern as the existing extractors.  Each dict must have:
       {"text": str, "label": str | None}
   where label is "ASSESSOR KEY", "ANSWER GUIDANCE", or None.
3. Add an elif branch in docx_to_text() to call your new extractor.
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.opc.exceptions import PackageNotFoundError
import json
import re
import tempfile
import os
import zipfile


# ===========================================================================
# Shared colour helpers
# ===========================================================================

def _is_red_rgb(rgb) -> bool:
    """Threshold used to classify an RGBColor as the answer-marking red: R≥200, G≤80, B≤80."""
    if rgb is None:
        return False
    return rgb[0] >= 200 and rgb[1] <= 80 and rgb[2] <= 80


def _is_red_run(run) -> bool:
    """
    Return True when a text run uses the red font colour used in Template A
    to mark correct answers.  Threshold: R≥200, G≤80, B≤80.
    """
    try:
        return _is_red_rgb(run.font.color.rgb)
    except Exception:
        return False


def _paragraph_style_is_red(para: Paragraph) -> bool:
    """
    Return True when the paragraph's own runs carry no direct colour, but
    the paragraph's NAMED STYLE defines red font colour instead (e.g. a
    custom "Answers" style with <w:color w:val="FF0000"/> in its style
    definition, applied via pStyle rather than per-run formatting).

    Some source documents are authored this way — a style is defined once
    and applied to every answer paragraph, instead of manually colouring
    each run red — so every run in the paragraph is colourless on its own
    and the red only exists in styles.xml. A run-only check (_is_red_run)
    misses every answer in a document built this way. Walks the style's
    base-style chain since a style may inherit its colour from a parent
    style rather than setting it directly, so this generalises to any
    future template that colours answers via a style, whatever that
    style happens to be named.
    """
    style = para.style
    while style is not None:
        try:
            rgb = style.font.color.rgb
        except Exception:
            rgb = None
        if rgb is not None:
            return _is_red_rgb(rgb)
        style = getattr(style, "base_style", None)
    return False


def _run_color_overridden(run) -> bool:
    """
    True when a run declares its own colour directly — including an
    explicit non-colour value like <w:color w:val="auto"/> — rather than
    silently inheriting whatever colour its paragraph style defines.

    "auto" has no RGB value (color.rgb is None, same as a run with no
    colour element at all), but color.type is AUTO instead of None, which
    is the only way to tell "this run opted out of the style's colour on
    purpose" apart from "this run never mentioned colour, so the style's
    colour applies." Needed because a document can mix the two within the
    same style: e.g. a reference-table header row using the same "Answers"
    style as the real red answers, but with each header run explicitly set
    to "auto" to render black instead of inheriting the style's red.
    """
    try:
        return run.font.color.type is not None
    except Exception:
        return False


def _para_red_info(para: Paragraph) -> tuple[str, bool]:
    """Return (stripped text, has_any_red_run) for a paragraph."""
    text = "".join(run.text for run in para.runs).strip()
    any_red = any(_is_red_run(r) and r.text.strip() for r in para.runs)
    if not any_red and text:
        # Only fall back to the paragraph's style colour when none of its
        # runs explicitly override colour themselves (red or otherwise) —
        # otherwise a paragraph that opts OUT of the style's red via an
        # explicit "auto" run colour (see _run_color_overridden) would be
        # wrongly treated as red just because it happens to share a style
        # with real answer paragraphs.
        no_overrides = not any(_run_color_overridden(r) for r in para.runs if r.text.strip())
        if no_overrides:
            any_red = _paragraph_style_is_red(para)
    return text, any_red


def _answer_label(text: str) -> str:
    """
    Choose the label prefix for a red-text item in Template A.

    Answer-guidance intro lines (e.g. 'Answer may address, but is not
    limited to, three of the following:') get a distinct ANSWER GUIDANCE
    tag so the AI extractor can separate them from the actual answer
    bullet points.

    All other red text → ASSESSOR KEY.
    """
    t = text.lower()
    if t.startswith("answer may") or t.startswith("answer must"):
        return "ANSWER GUIDANCE"
    return "ASSESSOR KEY"


# ===========================================================================
# List-numbering detection (a, b, c… option markers)
# ===========================================================================
#
# Word auto-numbers lettered MC options (a), b), c)…) via list-numbering
# XML (numPr -> numId -> numbering.xml -> abstractNum -> lvl -> numFmt).
# The letters themselves are NEVER present in the paragraph's run text —
# python-docx (and therefore the AI extractor) sees only the bare option
# text. A non-red option (an incorrect distractor) is then structurally
# indistinguishable from a stray short-answer bullet, and the AI has to
# guess the question type from wording alone. That guess fails whenever a
# multi-select MC question is phrased like an open recall question (e.g.
# "What are three (3) possible actions…") — it looks exactly like the
# short-answer example in the extractor's own system prompt.
#
# Fix: read numFmt directly from the document's numbering part and tag
# every paragraph in a lowerLetter/upperLetter list as an "OPTION" — a
# deterministic, colour- and wording-independent signal that this line is
# one of a fixed set of MC choices.

_LETTER_NUMFMTS = frozenset({"lowerLetter", "upperLetter"})
_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _build_numfmt_cache(doc: Document) -> dict[str, str]:
    """
    Maps numId -> numFmt ("lowerLetter", "bullet", "decimal", …) for level 0
    of every list definition in the document. Returns {} if the document
    has no numbering part (e.g. no auto-numbered lists at all).
    """
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        return {}
    if numbering_part is None:
        return {}

    numbering_elm = numbering_part.element

    abstract_fmt: dict[str, str] = {}
    for abstract_num in numbering_elm.findall(_W_NS + "abstractNum"):
        abstract_id = abstract_num.get(_W_NS + "abstractNumId")
        lvl0 = abstract_num.find(_W_NS + "lvl")
        if lvl0 is None:
            continue
        numfmt_el = lvl0.find(_W_NS + "numFmt")
        if numfmt_el is None:
            continue
        abstract_fmt[abstract_id] = numfmt_el.get(_W_NS + "val")

    numid_fmt: dict[str, str] = {}
    for num in numbering_elm.findall(_W_NS + "num"):
        num_id = num.get(_W_NS + "numId")
        abstract_id_el = num.find(_W_NS + "abstractNumId")
        if abstract_id_el is None:
            continue
        abstract_id = abstract_id_el.get(_W_NS + "val")
        if abstract_id in abstract_fmt:
            numid_fmt[num_id] = abstract_fmt[abstract_id]

    return numid_fmt


def _is_lettered_option(para: Paragraph, numfmt_cache: dict[str, str]) -> bool:
    """True when para belongs to a lowerLetter/upperLetter auto-numbered list."""
    p_pr = para._p.pPr
    if p_pr is None or p_pr.numPr is None:
        return False
    try:
        num_id = str(p_pr.numPr.numId.val)
    except Exception:
        return False
    return numfmt_cache.get(num_id) in _LETTER_NUMFMTS


# ===========================================================================
# Shared textbox helper (floating DrawingML text boxes)
# ===========================================================================

def _textbox_texts(para: Paragraph) -> list[tuple[str, bool]]:
    """
    Extract (text, is_red) pairs from any floating text boxes anchored
    to this paragraph element.  Handles DrawingML txbxContent elements.
    """
    out = []
    try:
        p_elm = para._p
    except Exception:
        return out

    for el in p_elm.iter():
        if not str(getattr(el, "tag", "")).endswith("}txbxContent"):
            continue
        for p2 in el.iter():
            if not str(getattr(p2, "tag", "")).endswith("}p"):
                continue
            parts: list[str] = []
            any_red = False
            for r in p2.iter():
                if not str(getattr(r, "tag", "")).endswith("}r"):
                    continue
                run_text = "".join(
                    t.text for t in r.iter()
                    if str(getattr(t, "tag", "")).endswith("}t")
                    and getattr(t, "text", None)
                )
                if run_text:
                    parts.append(run_text)
            text = "".join(parts).strip()
            if text:
                out.append((text, any_red))
    return out


# ===========================================================================
# Document block iterator (preserves body order)
# ===========================================================================

def _iter_blocks(doc: Document):
    """
    Yield top-level body children as Paragraph or Table objects in
    document order.  This preserves the visual sequence of content.
    """
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def _unique_row_cells(row):
    """
    Yield each cell in a table row once, skipping repeats caused by
    horizontally-merged cells.

    python-docx's row.cells returns the SAME underlying cell object once for
    every grid column a merged cell spans (a documented quirk, not a bug on
    our end) — a cell merged across 3 columns appears 3 times in row.cells,
    each with identical content. Left unguarded, a multi-paragraph merged
    cell (question text + several ASSESSOR KEY/OPTION lines) gets processed
    2-3x, since the existing "skip if identical to the immediately preceding
    item" dedup only catches single-line repeats, not a whole block replaying.
    """
    seen_tc_ids: set[int] = set()
    for cell in row.cells:
        tc_id = id(cell._tc)
        if tc_id in seen_tc_ids:
            continue
        seen_tc_ids.add(tc_id)
        yield cell


def _cell_plain_text(cell) -> str:
    """
    Resolve a cell's full text content as one plain string, recursing into
    a table nested inside the cell (rare double-nesting) by flattening it
    inline. Used only for the structured table representation handed to
    the AI extractor (see _table_to_dict) — unlike the paragraph-answer
    path, table cell content carries no ASSESSOR KEY / colour labelling
    here. The table is captured verbatim; deciding what's a header and
    what's the accepted answer is left entirely to the AI extractor.
    """
    parts: list[str] = []
    for block in cell.iter_inner_content():
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                parts.append(text)
        elif isinstance(block, Table):
            for row in block.rows:
                for sub_cell in _unique_row_cells(row):
                    sub_text = _cell_plain_text(sub_cell)
                    if sub_text:
                        parts.append(sub_text)
    return " ".join(parts)


def _table_to_dict(table) -> dict:
    """
    Render a table as a structured {"rows": [{"cells": [...]}]} dict
    instead of flattening it into inline text — preserves the table's
    exact row/column shape so the AI extractor receives the real
    structure rather than reconstructing it from a flattened stream.
    Deliberately does not try to detect or skip a header row here — that
    judgement, like which cell(s) hold the accepted answer, is left to the
    AI extractor, which has full context (question wording) this parsing
    layer doesn't.
    """
    rows: list[dict] = []
    for row in table.rows:
        cells = list(_unique_row_cells(row))
        if not cells:
            continue
        rows.append({"cells": [_cell_plain_text(c) for c in cells]})
    return {"rows": rows}


def _walk_cell(cell, process_para, push_table, on_new_question=None) -> None:
    """
    Walk a table cell's direct content in document order.

    cell.paragraphs never surfaces a table nested inside the cell (python-
    docx quirk, same family as the merged-cell repeat handled above), so
    without this a table-inside-a-cell — a reference grid (legislation /
    objective / regulator columns), a matching-question answer table, or a
    single-cell bordered answer box — is silently dropped from extraction.

    Plain paragraphs go through process_para exactly as before (unchanged
    colour / lettered-option labelling). A nested table is handled by shape:
      - 1 row x 1 column (just a bordered box wrapping free-response text,
        e.g. a plain answer box under a short-answer question) — no
        row/column structure to preserve, so its paragraph(s) are walked
        directly through process_para like any other cell text.
      - anything larger — captured as ONE structured table item via
        push_table(_table_to_dict(...)), keeping the whole grid intact for
        the AI extractor instead of flattening it into loose lines.

    on_new_question(), if given, is called right before process_para for a
    paragraph that immediately follows a nested table in this same cell.
    In Template A, every question's own text sits directly in the cell as
    plain paragraphs, and its answer content always lives inside a nested
    table (box or grid) — confirmed against the source — so that exact
    transition (table just ended, next block is a paragraph) IS the
    boundary between one question and the next. None (the default) is a
    no-op for callers that don't need this signal.
    """
    just_finished_table = False
    for block in cell.iter_inner_content():
        if isinstance(block, Paragraph):
            if just_finished_table and on_new_question is not None:
                on_new_question()
            process_para(block)
            just_finished_table = False
        elif isinstance(block, Table):
            rows = list(block.rows)
            n_cols = len(list(_unique_row_cells(rows[0]))) if rows else 0
            if len(rows) <= 1 and n_cols <= 1:
                for row in rows:
                    for sub_cell in _unique_row_cells(row):
                        _walk_cell(sub_cell, process_para, push_table, on_new_question)
            else:
                push_table(_table_to_dict(block))
            just_finished_table = True


# ===========================================================================
# Template detection
# ===========================================================================

def _detect_template(doc: Document) -> str:
    """
    Inspect the document structure and return a template identifier string.

    Detection strategy: look at the FIRST table's first cell text.
      "instructions"  → Template A (Learning Vault assessor-guide format)
      starts with "q" → Template B/C (per-question table format)
      anything else   → "unknown" (generic fallback)

    Body-level paragraphs before the table are intentionally ignored here
    because both Template A and B/C can have them.
    """
    for block in _iter_blocks(doc):
        if isinstance(block, Table):
            try:
                first_cell = block.rows[0].cells[0].text.strip().lower()
            except Exception:
                first_cell = ""

            if first_cell == "instructions":
                return "A"
            if first_cell.startswith("q"):
                return "B_C"
            # First table doesn't match known patterns — stop searching
            break

    return "unknown"


# ===========================================================================
# Template A — Learning Vault "Assessor Guide" single-table format
# ===========================================================================

# Column-0 labels that identify boilerplate / metadata rows in Template A.
# These rows must be skipped; only the question-content row(s) are extracted.
#
# To support a variant that adds new row labels (e.g. "Marking criteria"),
# extend this set — no other code changes are needed.
_TEMPLATE_A_SKIP_LABELS: frozenset[str] = frozenset({
    # ── Header boilerplate (top of table) ──────────────────────────────────
    "instructions",
    "range and conditions",
    "decision-making rules",
    "pre-approved reasonable adjustments",
    "rubric",
    # ── Footer boilerplate (bottom of table, SA format only) ───────────────
    "learner feedback",
    "assessment outcome",
    "assessor signature",
    "assessor name",
    "date",
    "final comments",
})


def _is_boilerplate_row_a(row) -> bool:
    """
    Return True when a Template A table row contains only boilerplate
    metadata that should not be included in the extracted text.

    Two conditions trigger a skip:
      1. The Col-0 label matches a known boilerplate name (e.g. "Rubric").
      2. All cells contain identical text that begins with "knowledge test"
         — these are section-heading rows that span the full table width.
    """
    try:
        col0 = row.cells[0].text.strip().lower()
    except Exception:
        return False

    # Condition 1: explicit boilerplate label in the first column
    if col0 in _TEMPLATE_A_SKIP_LABELS:
        return True

    # Condition 2: merged section-heading row such as "Knowledge Test 2 – MC"
    # All cells share the same text because the row is a full-width merge.
    unique_cell_texts = {c.text.strip().lower() for c in row.cells}
    if len(unique_cell_texts) == 1 and col0.startswith("knowledge test"):
        return True

    return False


def _extract_template_a(doc: Document) -> list[dict]:
    """
    Extract question content from a Template A document.

    Processing rules:
    • Boilerplate rows (header + footer) are skipped entirely — this
      eliminates false ASSESSOR KEY tags from the Rubric red bullets.
    • Within question-content rows, paragraphs are labelled:
        - "ANSWER GUIDANCE" for red intro lines like "Answer may address…"
        - "ASSESSOR KEY"    for red answer text / correct MC options
        - "OPTION"          for non-red lettered-list MC distractors
    • Floating text-box text is extracted if present.
    • Consecutive duplicate items (from merged cells) are deduplicated.
    """
    items: list[dict] = []
    last_key: tuple | None = None
    numfmt_cache = _build_numfmt_cache(doc)

    def push(text: str, label: str | None) -> None:
        nonlocal last_key
        text = text.strip()
        if not text:
            return
        key = (text, label)
        if key == last_key:
            return  # deduplicate merged-cell repetitions
        items.append({"text": text, "label": label})
        last_key = key

    def label_for_para(para: Paragraph) -> tuple[str, str | None]:
        text, is_red = _para_red_info(para)
        if is_red:
            return text, _answer_label(text)
        if _is_lettered_option(para, numfmt_cache):
            return text, "OPTION"
        return text, None

    def process_para(para: Paragraph) -> None:
        text, label = label_for_para(para)
        push(text, label)
        for tb_text, tb_red in _textbox_texts(para):
            push(tb_text, _answer_label(tb_text) if tb_red else None)

    def push_table(table_dict: dict) -> None:
        items.append({"table": table_dict})

    def after_answer_table() -> None:
        # A plain paragraph appearing right after a nested answer table
        # (see _walk_cell) is, structurally, the start of the NEXT
        # question — confirmed against the source: a question's own black
        # text sits directly in the cell, and its red-text answer content
        # is always inside a nested table. Once that nested table ends,
        # the next paragraph in the cell is a new question's stem, never
        # a continuation of the previous one.
        push("[QUESTION START]", None)

    for block in _iter_blocks(doc):
        if isinstance(block, Paragraph):
            # Body-level paragraphs in Template A are usually just the test
            # title (e.g. "Knowledge Test 2 - Multiple Choice") — keep them
            # as plain-text context for the AI extractor.
            text, _ = _para_red_info(block)
            push(text, None)

        elif isinstance(block, Table):
            for row in block.rows:
                if _is_boilerplate_row_a(row):
                    continue  # skip Instructions, Rubric, footer rows, etc.
                for cell in _unique_row_cells(row):
                    after_answer_table()  # marks this row/cell's own first question too
                    _walk_cell(cell, process_para, push_table, after_answer_table)

    return items


# ===========================================================================
# Template B / C — Per-question table format (SITHCCC / SITX style)
# ===========================================================================

# Substrings that identify body-level header paragraphs in Templates B/C.
# These appear before the question tables and contain student/admin metadata.
_TEMPLATE_BC_SKIP_SUBSTRINGS: tuple[str, ...] = (
    "knowledge test",        # e.g. "KNOWLEDGE TEST -- SHORT ANSWER"
    "student name",          # "Student Name: ___"
    "assessor:",             # "Assessor: ___"
    "result:",               # "Result: Satisfactory [ ]"
    "instructions:",         # "Instructions: Answer ALL questions"
    "write in full sentence", # "Write in full sentences where indicated."
    "sit30",                 # qualification code prefix e.g. "SIT30821"
    "total marks",           # "Total marks: 32"
)


def _is_header_para_bc(text: str) -> bool:
    """
    Return True when a body paragraph is metadata that should be skipped
    in Templates B/C (student admin header, qualification title, etc.).
    """
    t = text.lower()
    return any(pattern in t for pattern in _TEMPLATE_BC_SKIP_SUBSTRINGS)


def _extract_template_bc(doc: Document) -> list[dict]:
    """
    Extract question content from a Template B or C document.

    Processing rules:
    • Body-level metadata header paragraphs are skipped (student name,
      assessor, qualification title, etc.).
    • Each per-question table is extracted cell-by-cell in row order.
    • The "ASSESSOR KEY: " prefix is already embedded as literal text in
      the answer cell — no colour detection is required.
    • Consecutive duplicate items (from merged cells) are deduplicated.

    Template B (SA): last row of each table is a merged cell containing
      "ASSESSOR KEY: {full answer text}"

    Template C (MC): last row Col 2 contains "ASSESSOR KEY: {letter}"
      Options appear two per row across Col 1 and Col 2.
    """
    items: list[dict] = []
    last_text: str | None = None

    def push(text: str) -> None:
        nonlocal last_text
        text = text.strip()
        if not text:
            return
        if text == last_text:
            return  # deduplicate merged-cell repetitions
        items.append({"text": text, "label": None})
        last_text = text

    def handle_para(para: Paragraph) -> None:
        push(para.text)
        for tb_text, _ in _textbox_texts(para):
            push(tb_text)

    def push_table(table_dict: dict) -> None:
        items.append({"table": table_dict})

    for block in _iter_blocks(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text and not _is_header_para_bc(text):
                push(text)

        elif isinstance(block, Table):
            for row in block.rows:
                for cell in _unique_row_cells(row):
                    _walk_cell(cell, handle_para, push_table)

    return items


# ===========================================================================
# Generic fallback (unknown / future template)
# ===========================================================================

def _extract_generic(doc: Document) -> list[dict]:
    """
    Fallback extractor for unrecognised templates.

    Extracts all text in document order using red-text detection for
    answer labelling.  No boilerplate filtering is applied.

    This is a best-effort path.  If a new template type is encountered
    regularly, implement a dedicated extractor (see module docstring).
    """
    items: list[dict] = []
    last_key: tuple | None = None

    def push(text: str, label: str | None) -> None:
        nonlocal last_key
        text = text.strip()
        if not text:
            return
        key = (text, label)
        if key == last_key:
            return
        items.append({"text": text, "label": label})
        last_key = key

    def label_for_para(para: Paragraph) -> tuple[str, str | None]:
        text, is_red = _para_red_info(para)
        return text, (_answer_label(text) if is_red else None)

    def process_para(para: Paragraph) -> None:
        text, label = label_for_para(para)
        push(text, label)
        for tb_text, tb_red in _textbox_texts(para):
            push(tb_text, _answer_label(tb_text) if tb_red else None)

    def push_table(table_dict: dict) -> None:
        items.append({"table": table_dict})

    for block in _iter_blocks(doc):
        if isinstance(block, Paragraph):
            process_para(block)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in _unique_row_cells(row):
                    _walk_cell(cell, process_para, push_table)

    return items


# ===========================================================================
# Private helper: instruction extraction from an already-loaded Document
# ===========================================================================

def _instructions_from_doc(doc: Document, template: str) -> str:
    """
    Extract the boilerplate instruction block from an already-loaded Document.

    Called by extract_document() on the same doc object used for raw_text,
    so the file is only read and parsed once.

    Template A only — returns "" for all other templates.
    """
    if template != "A":
        return ""

    sections: list[str] = []
    seen: set[str] = set()

    for block in _iter_blocks(doc):
        if not isinstance(block, Table):
            continue

        for row in block.rows:
            if not _is_boilerplate_row_a(row):
                continue  # skip question rows

            col0 = row.cells[0].text.strip()

            # Skip merged section-heading rows ("Knowledge Test N …")
            if col0.lower().startswith("knowledge test"):
                continue

            # Pick the content cell by content length rather than a fixed
            # column index — robust to columns being added, removed, or
            # reordered. Column 0 is always the label ("Instructions",
            # "Rubric", …); the content lives in whichever remaining cell
            # has the most text (for 3-column SA tables, columns 1 and 2
            # are merged duplicates of the same text — length-based
            # selection handles that too, since either one wins equally).
            candidates  = row.cells[1:] or row.cells[:1]
            content_cell = max(candidates, key=lambda c: len(c.text))
            content_lines = [
                p.text.strip()
                for p in content_cell.paragraphs
                if p.text.strip()
            ]
            content = "\n".join(content_lines)

            # Deduplicate identical sections that appear across merged cells
            key = f"{col0}|{content}"
            if key in seen:
                continue
            seen.add(key)

            if content:
                sections.append(f"[{col0}]\n{content}")

    return "\n\n".join(sections)


# ===========================================================================
# Points — global per-question mark value stated in the instructions block
# ===========================================================================

_POINTS_PATTERN = re.compile(r"graded out of\s+(\d+(?:\.\d+)?)\s*marks?", re.IGNORECASE)


def parse_points_per_question(instructions_text: str) -> float | None:
    """
    Extract the global per-question mark value from the Instructions block
    (Template A only), e.g. "Each question is graded out of 1 mark." →  1.0.

    This is the authoritative source for points when present — every sample
    Template A document states this uniformly, so it should override
    whatever value the AI extractor guessed per-question. Returns None when
    no such statement is found; callers should fall back to the
    per-question AI-extracted value in that case, and leave points unset
    if that's also absent rather than inventing a number.
    """
    if not instructions_text:
        return None
    match = _POINTS_PATTERN.search(instructions_text)
    if not match:
        return None
    return float(match.group(1))


# ===========================================================================
# Public API
# ===========================================================================

def extract_document(file_bytes: bytes, filename: str) -> tuple[str, str]:
    """
    Parse a .docx file ONCE and return both extracted texts as a tuple:
        (raw_text, instructions_text)

    raw_text        — question content only, with ASSESSOR KEY / ANSWER GUIDANCE
                      labels, ready for the AI extractor.
    instructions_text — boilerplate section (Instructions, Rubric, Range, etc.)
                        stored as ground truth for future validation.
                        Empty string for Templates B/C which have no structured
                        instruction block.

    The file is opened exactly once; both values are derived from the same
    Document object so there is no redundant disk or parse overhead.
    """
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        try:
            doc = Document(tmp_path)
        except (zipfile.BadZipFile, KeyError, PackageNotFoundError) as e:
            # PackageNotFoundError/BadZipFile: the file isn't a zip archive
            # at all (e.g. a .pdf, .txt, or image renamed to .docx).
            # KeyError: it IS a zip but missing required Word package parts
            # (e.g. a .xlsx/.pptx renamed to .docx, or a corrupted file).
            raise ValueError(
                f"'{filename}' doesn't appear to be a valid .docx file "
                f"(it may be a different format, renamed, or corrupted): {e}"
            ) from e

        template = _detect_template(doc)

        # --- Question content (raw_text) ------------------------------------
        if template == "A":
            items = _extract_template_a(doc)
        elif template == "B_C":
            items = _extract_template_bc(doc)
        else:
            items = _extract_generic(doc)

        lines: list[str] = []
        for item in items:
            if "table" in item:
                # A structured table nested inside a question's cell (see
                # _table_to_dict) — emitted as one JSON blob rather than
                # flattened text, so the AI extractor gets the table's real
                # row/column shape instead of reconstructing it from a
                # flattened stream.
                lines.append("TABLE_DATA: " + json.dumps(item["table"], ensure_ascii=False))
                lines.append("")
                continue
            label = item.get("label")
            text  = item["text"]
            lines.append(f"{label}: {text}" if label else text)
            lines.append("")  # blank separator between items

        raw_text = "\n".join(lines).strip()

        # --- Instruction block (instructions_text) --------------------------
        # Reuses the same doc object — no second file read needed.
        instructions_text = _instructions_from_doc(doc, template)

        return raw_text, instructions_text

    finally:
        os.unlink(tmp_path)


def docx_to_text(file_bytes: bytes, filename: str) -> str:
    """
    Convenience wrapper — returns only raw_text.
    Prefer extract_document() when you also need instructions_text.
    """
    raw_text, _ = extract_document(file_bytes, filename)
    return raw_text
